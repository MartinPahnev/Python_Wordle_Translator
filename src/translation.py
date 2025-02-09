"""
This module handles the translation of text and files. It supports translating 
plain text and DOCX files with language detection and Google Translate API.
"""

import os
import io
from flask import render_template, request, flash, send_file
from docx import Document
from langdetect import detect

def translate_text(request, translator):
    """Translate text from a source language to a target language."""
    if request.method == 'POST':
        text = request.form.get('text')
        target_lang = request.form.get('target_lang', 'en')
        if not text:
            flash("Моля, въведете текст за превод.")
            return redirect(request.url)
        try:
            source_lang = detect(text)
        except Exception:
            source_lang = 'auto'
        translated = translator.translate(text, src=source_lang, dest=target_lang)
        return render_template('translate.html',
                               original=text,
                               translated=translated.text,
                               source_lang=source_lang,
                               target_lang=target_lang)
    return render_template('translate.html')

def translate_file(request, translator):
    """Translate text from a file and return the translated file."""
    if request.method == 'POST':
        target_lang = request.form.get('target_lang', 'en')
        file = request.files.get('file')
        if not file:
            flash("Няма качен файл!")
            return redirect(request.url)
        filename = file.filename
        ext = os.path.splitext(filename)[1].lower()
        content = ""
        if ext == '.txt':
            content = file.read().decode('utf-8')
        elif ext == '.docx':
            doc = Document(file)
            content = "\n".join([p.text for p in doc.paragraphs])
        else:
            flash("Неподдържан тип файл!")
            return redirect(request.url)
        try:
            source_lang = detect(content)
        except Exception:
            source_lang = 'auto'
        translated = translator.translate(content, src=source_lang, dest=target_lang)
        output = io.BytesIO()
        if ext == '.txt':
            output.write(translated.text.encode('utf-8'))
            output.seek(0)
            return send_file(output, as_attachment=True, download_name="translated.txt", mimetype='text/plain')
        elif ext == '.docx':
            new_doc = Document()
            new_doc.add_paragraph(translated.text)
            new_output = io.BytesIO()
            new_doc.save(new_output)
            new_output.seek(0)
            return send_file(new_output, as_attachment=True, download_name="translated.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return render_template('upload.html')
